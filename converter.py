import sys
import csv
from docx import Document
from docx.shared import Inches
import argparse

parser = argparse.ArgumentParser(description="converter.py")
parser.add_argument("csv_filename")
parser.add_argument("--generate-numbering", help="generate numbering instead of using csv numbering", action="store_true")

args = parser.parse_args()

document = Document()
document.add_heading("User Stories")

table = document.add_table(rows=1, cols=4)
table.style = "Table Grid"
header_rows = table.rows[0].cells
header_rows[0].text = "#"
header_rows[1].text = "Story"
header_rows[2].text = "Acceptance criteria"
header_rows[3].text = "Pts"

# def show_help():
#     print("Usage: converter.py <stories.csv>")
        

# if "-h" in sys.argv:
#     show_help()
#     sys.exit(0)

# elif len(sys.argv) != 2:
#     show_help()
#     sys.exit(1)


# infile_name = sys.argv[1]
infile_name = args.csv_filename
generate_numbering = args.generate_numbering

with open(infile_name) as f:
    stories_reader = csv.reader(f)
    for i, row in enumerate(stories_reader):

        # Don't include csv header
        if i==0:
            continue

        tag, num, story, points, add_info, excluded = row
        acc_criteria = add_info.split("\n")

        row_cells = table.add_row().cells
        row_cells[0].text = str(num[len("STORY#"):]) if not generate_numbering else str(i)
        row_cells[1].text = str(story)
        # row_cells[2].text = add_info
        for a in acc_criteria:
            if len(a) < 2:
                continue
            # pg.add_run(a)
            pg = row_cells[2].add_paragraph(a, style="List Bullet")
        row_cells[3].text = str(points)


if infile_name[-4:] == ".csv":
    outfile_name = infile_name[:-4] + ".docx"
else:
    outfile_name = infile_name + ".docx"

document.save(outfile_name)
print("Written to " + outfile_name)